import os
from django.conf import settings
from datetime import datetime
from docx import Document
import pandas as pd

def ghi_du_lieu_table(document, data):
    for k, v in data.items():
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        inline = paragraph.runs
                        for j in range(len(inline)):
                            if k in inline[j].text:
                                inline[j].text = inline[j].text.replace(
                                    k, v)

def ghi_du_lieu_para(document, data):
    for k, v in data.items():
        for paragraph in document.paragraphs:
            inline = paragraph.runs
            for j in range(len(inline)):
                if k in inline[j].text:
                    inline[j].text = inline[j].text.replace(k, v)

def del_row(table, row_start_del, row_count):
    # Xóa các hàng không có dữ liệu
    for row in table.rows[row_start_del + 1:row_count]:
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

def ghi_du_lieu_cell(row, text, text_replace):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                inline = paragraph.runs
                for j in range(len(inline)):
                    if text in inline[j].text:
                        inline[j].text = inline[j].text.replace(
                            text, text_replace)

def vi_tri_bang(document, repalce_text):
    # document = Document(path)
    i = 0
    for paragraph in document.paragraphs:
        inline = paragraph.runs
        for j in range(len(inline)):
            if repalce_text in inline[j].text:
                return i
        i = i + 1

class lap_qd_ktra:
    def __init__(self, tt_qd, doan_ktra):
        self.tt_qd = tt_qd
        self.doan_ktra = doan_ktra

    def to_trinh(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "1.to_trinh_ktr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_To_trinh.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)   
        document.save(path)
        return path
    
    def qd_gsat(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "1.qd_giam_sat_ktr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_QD_giam_sat.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)   
        document.save(path)
        return path

    def kh_gsat(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.kh_giam_sat.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "1.kh_giam_sat_ktr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_KH_giam_sat.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path

    def qd_ktra(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "1.qd_ktra.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        # ghi dữ liệu thành phần đoàn
        doc_table = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "doan_ktra_table.docx"))
        table = doc_table.tables[0]
        # del_row(table, len(doan_ktra['<ten_cb>']), len(table.rows))
        for i in range(len(self.doan_ktra['<ten_cb>'])):
            row = table.rows[i]
            for k,v in self.doan_ktra.items():
                ghi_du_lieu_cell(row, k, v[i])
        # Xóa các hàng không có dữ liệu
        del_row(table, len(self.doan_ktra['<ten_cb>']) - 1, len(table.rows))
        # doc_table.save("doan_ktra.docx")
        # tb = Document("doan_ktra.docx")
        # template = tb.tables[0]
        tbl = table._tbl
        # tbl = template._tbl
        # Tìm vị trí '[*]' (vị trí đánh dấu)
        # Copy bảng thành viên vào vị trí đánh dấu
        loc = vi_tri_bang(document, '[*]')
        # template = tb.tables[0]
        # tbl = template._tbl
        p = document.paragraphs[loc]._p
        p.addnext(tbl)
        filename = self.tt_qd["<mst>"] + "_QD_kiem_tra.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path
    
    def empty_media(self):
        media_store = os.path.join(settings.STATICFILES_DIRS[0], "media_store")
        for file in os.listdir(media_store):
            path = os.path.join(media_store, file) 
            os.remove(path)

class lap_qd_ttra:
    def __init__(self, tt_qd, doan_ttra):
        self.tt_qd = tt_qd
        self.doan_ttra = doan_ttra

    def to_trinh(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.to_trinh_ttr.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "2.to_trinh_ttr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_To_trinh.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)   
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)      
        document.save(path)
        return path
    
    def qd_gsat(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.qd_giam_sat _ttr.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "2.qd_giam_sat_ttr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_QD_giam_sat.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path
    
    def kh_gsat(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.kh_giam_sat.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "2.kh_giam_sat_ttr.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_KH_giam_sat.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path

    def qd_ttra(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.qd_ttra.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "2.qd_ttra.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        # ghi dữ liệu thành phần đoàn
        table = document.tables[0]
        # del_row(table, len(doan_ktra['<ten_cb>']), len(table.rows))
        for i in range(len(self.doan_ttra['<ten_cb>'])):
            row = table.rows[i]
            for k,v in self.doan_ttra.items():
                ghi_du_lieu_cell(row, k, v[i])
        # Xóa các hàng không có dữ liệu
        del_row(table, len(self.doan_ttra['<ten_cb>']) - 1, len(table.rows))
        filename = self.tt_qd["<mst>"] + "_QD_thanh_tra.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path
    
    def kh_ttra(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.kh_ttra.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "2.kh_ttra.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        table = document.tables[1]
        # del_row(table, len(doan_ktra['<ten_cb>']), len(table.rows))
        for i in range(len(self.doan_ttra['<ten_cb>'])):
            row = table.rows[i]
            for k,v in self.doan_ttra.items():
                ghi_du_lieu_cell(row, k, v[i])
        # Xóa các hàng không có dữ liệu
        del_row(table, len(self.doan_ttra['<ten_cb>']) - 1, len(table.rows))
        filename = self.tt_qd["<mst>"] + "_KH_thanh_tra.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path
    
    def empty_media(self):
        media_store = os.path.join(settings.STATICFILES_DIRS[0], "media_store")
        for file in os.listdir(media_store):
            path = os.path.join(media_store, file)
            os.remove(path)

class lap_qd_ktra_hoan_gtgt:
    def __init__(self, tt_qd, doan_ktra):
            self.tt_qd = tt_qd
            self.doan_ktra = doan_ktra

    def to_trinh(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "3.to_trinh_ktr_hoan_gtgt.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_To_trinh.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)   
        document.save(path)
        return path

    def qd_gsat(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "3.qd_giam_sat_hoan_gtgt.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_QD_giam_sat.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)   
        document.save(path)
        return path

    def kh_gsat(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.kh_giam_sat.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "3.kh_giam_sat_hoan_gtgt.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_KH_giam_sat.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path
    
    def qd_ktra(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.qd_ttra.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "3.qd_ktra_hoan_gtgt.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        # ghi dữ liệu thành phần đoàn
        table = document.tables[0]
        # del_row(table, len(doan_ktra['<ten_cb>']), len(table.rows))
        for i in range(len(self.doan_ktra['<ten_cb>'])):
            row = table.rows[i]
            for k,v in self.doan_ktra.items():
                ghi_du_lieu_cell(row, k, v[i])
        # Xóa các hàng không có dữ liệu
        del_row(table, len(self.doan_ktra['<ten_cb>']) - 1, len(table.rows))
        filename = self.tt_qd["<mst>"] + "_QD_kiem_tra.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path

    def empty_media(self):
        media_store = os.path.join(settings.STATICFILES_DIRS[0], "media_store")
        for file in os.listdir(media_store):
            path = os.path.join(media_store, file) 
            os.remove(path)

class lap_qd_ktra_giai_the:
    def __init__(self, tt_qd, doan_ktra):
            self.tt_qd = tt_qd
            self.doan_ktra = doan_ktra
    
    def to_trinh(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "4.to_trinh_ktr_giai_the.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_To_trinh.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)   
        document.save(path)
        return path

    def qd_gsat(self):
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "4.qd_giam_sat_ktr_giai_the.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_QD_giam_sat.docx"
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)   
        document.save(path)
        return path

    def kh_gsat(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.kh_giam_sat.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "4.kh_giam_sat_ktr_giai_the.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        filename = self.tt_qd["<mst>"] + "_KH_giam_sat.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path

    def qd_ktra(self):
        # document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media/2.qd_ttra.docx"))
        document = Document(os.path.join(settings.STATICFILES_DIRS[0], "media", "4.qd_ktra_giai_the.docx"))
        ghi_du_lieu_para(document, self.tt_qd)
        # ghi dữ liệu thành phần đoàn
        table = document.tables[0]
        # del_row(table, len(doan_ktra['<ten_cb>']), len(table.rows))
        for i in range(len(self.doan_ktra['<ten_cb>'])):
            row = table.rows[i]
            for k,v in self.doan_ktra.items():
                ghi_du_lieu_cell(row, k, v[i])
        # Xóa các hàng không có dữ liệu
        del_row(table, len(self.doan_ktra['<ten_cb>']) - 1, len(table.rows))
        filename = self.tt_qd["<mst>"] + "_QD_kiem_tra.docx"
        # path = os.path.join(settings.STATICFILES_DIRS[0], "media_store/" + filename)
        path = os.path.join(settings.STATICFILES_DIRS[0], "media_store", filename)
        document.save(path)
        return path

    def empty_media(self):
        media_store = os.path.join(settings.STATICFILES_DIRS[0], "media_store")
        for file in os.listdir(media_store):
            path = os.path.join(media_store, file) 
            os.remove(path)
    
# <ngay_thang>, <phieu_xly_ngay>, <sl_cb>, <cb_cv>, <so_ngay_ktra>, <ngay_ktra>
# <ten_dv>, <mst>, <dia_chi>, <ky_hoan_thue>
# <hinh_thuc_ky>, <LD_CUC>, <ld_cuc_ten>, <LD_PHONG>, <ld_phong_ten> 
""" tt_qd = {
    '<qd_tkt_tct>': 'Quyết định số 2271/QĐ-TCT ngày 30 tháng 12 năm 2020', #qd ttr
    '<nam_qd_tkt_tct>': '2020',
    '<ngay_thang>' : "ngày      tháng 8 năm 2021",
    '<ten_dv>' : 'Công ty TNHH ABC',
    '<mst>' : '3200123456',
    '<dia_chi>' : '260 Hùng Vương, tp Đông Hà, Quảng Trị',
    "<sl_cb>" : '03',
    '<truong_doan_ttr>': 'Phạm Văn Vui',
    "<cb_cv>" : 'Phó trưởng phòng',
    '<so_nam_ktra>' : '01',
    '<nam_ktra>' : '2020',
    '<so_ngay_ktra>' : '05',         
    '<ngay_ktra>' : "ngày 06 tháng 9 năm 2020",
    '<ng_giam_sat>' : 'ông Nguyễn Tiền Hải',
    '<Ng_giam_sat>' : 'Ông Nguyễn Tiền Hải',
    '<ng_giam_sat_cv>' : 'Phó trưởng phòng', 
    '<ld_phong>' : 'Phó trưởng phòng',
    '<LD_PHONG>': 'PHÓ TRƯỞNG PHÒNG',
    '<ld_phong_ten>' : 'Nguyễn Tiền Hải',
    '<ld_cuc>' : 'PHÓ CỤC TRƯỞNG',
    '<ld_cuc_ten>' : 'Dương Quốc Hoàn',
    '<hinh_thuc_ky>' : 'KT.CỤC TRƯỞNG',
    '<noi_nhan>': 'CCT KV Vĩnh Linh - Gio Linh',
}
doan_ttra = {
    "<ten_cb>" : ['Ông: Phạm Văn Vui', 'Bà: Nguyễn Thị Thanh Huyền', "Bà: Hà Thị Thanh Thủy"],
    "<ngach_cb>" : ['Kiểm soát viên thuế', 'Kiểm tra viên thuế', 'Kiểm soát viên thuế'],
    "<cv_doan>" : ['Trưởng đoàn', 'Thành viên', 'Thành viên'],
} """

tt_qd = {
    "<ngay_thang>": "ngày      tháng 3 năm 2021", # nhập tháng
    "<phieu_xly_ngay>": "22/6/2021",
    # "<hs_hoan_so>": "01", 
    # "<hs_hoan_ngay>": "01/10/2021", 
    # "<hoan_tien>": "130.032.903.522",
    # "<ky_hoan_thue>": "tháng 2/2020 đến tháng 9/2021",
    # "<so_qd>": "2271/QĐ-TCT của Tổng cục Thuế về việc phê duyệt kế hoạch thanh tra, kiểm tra thuế tại doanh nghiệp năm 2021", # nhập số QĐ
    "<ten_dv>" : "Công Ty TNHH Xây Dựng Thủy Điện Đakrông", # lấy từ dữ liệu
    "<dia_diem_ktra>": "tại trụ sở của cơ quan thuế",
    "<mst>" : "3200172428", # nhập
    "<dia_chi>" : "Xã Gio Châu, huyện Gio Linh, tỉnh Quảng Trị", # lấy từ dữ liệu  
    "<sl_cb>" : "03", # đếm số lượng cán bộ
    "<cb_cv>" : "Phạm Văn Vui – Phó trưởng phòng", # giới tính : tên - chức vụ (nếu có), ktra giải thể: đồng chí
    # "<so_nam_ktra>" : "01", # nhập
    "<nam_ktra>" : "2018", # tùy số lượng năm kiểm tra
    "<so_ngay_ktra>" : "01", # nhập
    "<ngay_ktra>" : "ngày 02 tháng 9 năm 2021", # nhập
    "<ng_giam_sat>" : "ông Nguyễn Tiền Hải", # nhập
    '<Ng_giam_sat>' : "Ông Nguyễn Tiền Hải",
    "<ng_giam_sat_cv>" : "Phó Trưởng phòng", # lấy từ dữ liệu
    "<LD_PHONG>" : "phó trưởng phòng".upper(), # chọn ['phó trưởng phòng', 'trưởng phòng']
    '<ld_phong>' : "Phó trưởng phòng",
    "<ld_phong_ten>" : "nguyễn tiền hải".title (), # nhập
    "<hinh_thuc_ky>" : "KT.CỤC TRƯỞNG", # chọn
    "<LD_CUC>" : "PHÓ CỤC TRƯỞNG", 
    "<ld_cuc_ten>" : "dương quốc hoàn".title(), # nhập
}
doan_ktra = {
    "<ten_cb>" : ['Ông: Phạm Văn Vui', 'Bà: Nguyễn Thị Thanh Huyền', "Bà: Hà Thị Thanh Thủy"],
    "<cv_cb>" : ['Phó trưởng phòng', 'Kiểm tra viên chính thuế', 'Kiểm soát viên thuế'],
    "<cv_doan>" : ['Trưởng đoàn', 'Thành viên', 'Thành viên'],
}
class cap_nhat_nnt:
    def __init__(self, path):
        self.path = path
        self.df = pd.read_excel(self.path, converters={'Mã số thuế':str})

    def check_data_valid(self):
        df = self.df
        errors = []
        for row in range(df.shape[0]):
            if pd.isnull(df.iloc[row, :]).any():
                errors.append(f"Điền đầy đủ thông tin NNT tại dòng dữ liệu thứ {row+1}")
        return errors if errors else True

df = pd.read_excel("D:\\python\\demo-tkt\\static\\media\\cap_nhat_nnt.xlsx", converters={'Mã số thuế':str})
""" for row in range(df.shape[0]):
    print(df.iloc[row]['Mã số thuế']) """
        
    
